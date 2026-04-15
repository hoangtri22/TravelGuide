# -*- coding: utf-8 -*-
"""AdminWeb_MaNguon_HamVaFile.docx — chỉ các hàm/route dùng trong sequence diagram Admin."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = Path(__file__).resolve().parent.parent / "AdminWeb_MaNguon_HamVaFile.docx"
ROOT = Path(__file__).resolve().parent.parent


def add_table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            t.rows[ri].cells[ci].text = str(val)
    doc.add_paragraph()


def main():
    doc = Document()
    t = doc.add_paragraph()
    t.add_run(
        "Ánh xạ hàm / file — theo Sequence Diagram Admin"
    ).bold = True
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Project: {ROOT}")
    doc.add_paragraph(
        "Chỉ gồm luồng: Đăng nhập, loadPois, Lưu POI, Xóa POI, "
        "Duyệt/Từ chối, Xuất JSON, Bản dịch, Tài khoản."
    )

    doc.add_heading("Phía trình duyệt — wwwroot/app.js", level=1)
    add_table(
        doc,
        ["Thành phần (sequence)", "File", "Tên trong code"],
        [
            ("Gọi HTTP + Bearer", "wwwroot/app.js", "api(url, options)"),
            ("Đăng nhập submit", "wwwroot/app.js", 'loginForm addEventListener("submit")'),
            ("Sau login: tải POI", "wwwroot/app.js", "loadPois()"),
            ("Sau login: tài khoản", "wwwroot/app.js", "loadAccounts()"),
            ("Sau login: bản dịch POI đầu", "wwwroot/app.js", "loadTranslation(pois[0].id)"),
            ("Lưu / tạo POI", "wwwroot/app.js", 'poiForm addEventListener("submit")'),
            ("Xóa / Duyệt / Từ chối", "wwwroot/app.js", 'poiTable addEventListener("click")'),
            ("Đổi POI bản dịch", "wwwroot/app.js", 'translationPoiSelect addEventListener("change")'),
            ("Lưu bản dịch", "wwwroot/app.js", 'translationForm addEventListener("submit")'),
            ("Tạo tài khoản", "wwwroot/app.js", 'accountForm addEventListener("submit")'),
            ("Xóa tài khoản", "wwwroot/app.js", 'accountTable addEventListener("click")'),
            ("Xuất file", "wwwroot/app.js", 'exportBtn addEventListener("click")'),
            ("Điền form POI (Sửa)", "wwwroot/app.js", "fillPoiForm(p)"),
            ("Reset form POI", "wwwroot/app.js", "clearPoiForm()"),
        ],
    )

    doc.add_heading("Phía server — TravelGuide.AdminWeb/Program.cs", level=1)

    doc.add_paragraph("Xác thực (mọi request có Bearer):")
    add_table(
        doc,
        ["Hàm", "File"],
        [
            ("Authenticate(HttpContext context, AuthStore authStore)", "Program.cs"),
            ("AuthStore.GetPrincipal(string token)", "Program.cs"),
            ("AuthStore.CreateToken(UserAccount user)", "Program.cs"),
        ],
    )

    doc.add_paragraph("Luồng 1 — Đăng nhập:")
    add_table(
        doc,
        ["Bước", "File", "Hành động trong code"],
        [
            ("POST login", "Program.cs", "app.MapPost(\"/api/auth/login\", ...)"),
            ("", "Program.cs", "TravelGuideDb.GetUserByUsernameAsync(request.Username)"),
            ("", "Program.cs", "PasswordTools.Hash(request.Password)"),
            ("", "Program.cs", "authStore.CreateToken(user)"),
        ],
    )

    doc.add_paragraph("Luồng 2 — GET /api/pois (loadPois):")
    add_table(
        doc,
        ["Bước", "File", "Hành động trong code"],
        [
            ("", "Program.cs", "app.MapGet(\"/api/pois\", ...)"),
            ("", "Program.cs", "TravelGuideDb.EnsureAutoTranslationsAsync(httpClient, targetLang: lang)"),
            ("", "Program.cs", "FillLangAsync(...) → MyMemoryTranslator.TranslateAsync(...)"),
            ("", "Program.cs", "TravelGuideDb.UpdateTranslationsAsync(...) (khi có thay đổi)"),
            ("", "Program.cs", "TravelGuideDb.GetPoisAsync(includeUnpublished: true) // admin"),
        ],
    )

    doc.add_paragraph("Luồng 3 — POST/PUT /api/pois (Lưu POI, gồm audioUrl trong PoiDto):")
    add_table(
        doc,
        ["Bước", "File", "Hành động trong code"],
        [
            ("POST", "Program.cs", "app.MapPost(\"/api/pois\", ...)"),
            ("", "Program.cs", "TravelGuideDb.CreatePoiAsync(poi, principal)"),
            ("", "Program.cs", "TravelGuideDb.EnsureAutoTranslationsAsync(httpClient, id)"),
            ("PUT", "Program.cs", "app.MapPut(\"/api/pois/{id:int}\", ...)"),
            ("", "Program.cs", "TravelGuideDb.UpdatePoiAsync(id, poi, principal)"),
            ("", "Program.cs", "TravelGuideDb.EnsureAutoTranslationsAsync(httpClient, id)"),
            ("Bind field POI", "Program.cs", "BindPoi(SqliteCommand, PoiDto) // gồm AudioUrl"),
        ],
    )

    doc.add_paragraph("Luồng 4 — DELETE /api/pois/{id}:")
    add_table(
        doc,
        ["Bước", "File", "Hành động trong code"],
        [
            ("", "Program.cs", "app.MapDelete(\"/api/pois/{id:int}\", ...)"),
            ("", "Program.cs", "TravelGuideDb.DeletePoiAsync(id)"),
        ],
    )

    doc.add_paragraph("Luồng 5 — Duyệt / Từ chối:")
    add_table(
        doc,
        ["Bước", "File", "Hành động trong code"],
        [
            ("", "Program.cs", "app.MapPut(\"/api/pois/{id:int}/approve\", ...)"),
            ("", "Program.cs", "TravelGuideDb.SetPoiStatusAsync(id, \"published\", \"\")"),
            ("", "Program.cs", "app.MapPut(\"/api/pois/{id:int}/reject\", ...)"),
            ("", "Program.cs", "RejectPoiRequest → SetPoiStatusAsync(id, \"rejected\", reason)"),
        ],
    )

    doc.add_paragraph("Luồng 6 — Xuất extra_places.json:")
    add_table(
        doc,
        ["Bước", "File", "Hành động trong code"],
        [
            ("", "Program.cs", "app.MapGet(\"/api/export/extra_places.json\", ...)"),
            ("", "Program.cs", "TravelGuideDb.GetExportPlacesAsync()"),
        ],
    )

    doc.add_paragraph("Luồng 7 — Bản dịch:")
    add_table(
        doc,
        ["Bước", "File", "Hành động trong code"],
        [
            ("GET", "Program.cs", "app.MapGet(\"/api/translations/{id:int}\", ...)"),
            ("", "Program.cs", "EnsureAutoTranslationsAsync(httpClient, id, lang)"),
            ("", "Program.cs", "TravelGuideDb.GetPoiAsync(id)"),
            ("PUT", "Program.cs", "app.MapPut(\"/api/translations/{id:int}\", ...)"),
            ("", "Program.cs", "TravelGuideDb.UpdateTranslationsAsync(id, TranslationUpdateRequest)"),
        ],
    )

    doc.add_paragraph("Luồng 8 — Tài khoản:")
    add_table(
        doc,
        ["Bước", "File", "Hành động trong code"],
        [
            ("GET", "Program.cs", "app.MapGet(\"/api/accounts\", ...)"),
            ("", "Program.cs", "TravelGuideDb.GetUsersAsync()"),
            ("POST", "Program.cs", "app.MapPost(\"/api/accounts\", ...)"),
            ("", "Program.cs", "TravelGuideDb.CreateUserAsync(...) // dùng PasswordTools.Hash bên trong"),
            ("DELETE", "Program.cs", "app.MapDelete(\"/api/accounts/{id:int}\", ...)"),
            ("", "Program.cs", "TravelGuideDb.DeleteUserAsync(id)"),
        ],
    )

    doc.add_heading("DTO / record xuất hiện trong các luồng trên", level=1)
    add_table(
        doc,
        ["Record", "File"],
        [
            ("LoginRequest", "Program.cs"),
            ("PoiDto", "Program.cs"),
            ("TranslationUpdateRequest", "Program.cs"),
            ("RejectPoiRequest", "Program.cs"),
            ("CreateAccountRequest", "Program.cs"),
            ("UserAccount", "Program.cs"),
            ("AuthPrincipal", "Program.cs"),
            ("ExportPoi", "Program.cs"),
        ],
    )

    doc.save(OUT)
    print(f"Wrote: {OUT}")


if __name__ == "__main__":
    main()
