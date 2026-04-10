# -*- coding: utf-8 -*-
"""
Tạo Word: danh sách file nguồn TravelGuide + AdminWeb và mô tả tổng quan (tiếng Việt).

Chạy (từ repo):
  pip install python-docx
  python TravelGuide.AdminWeb/scripts/build_file_overview_docx.py

Output: TravelGuide_TongQuan_File.docx (thư mục gốc dự án TravelGuide).
"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE = Path(__file__).resolve().parent
ADMIN_WEB = HERE.parent
REPO_ROOT = ADMIN_WEB.parent
OUT = REPO_ROOT / "TravelGuide_TongQuan_File.docx"

# (đường dẫn tương đối từ REPO_ROOT, mô tả ngắn)
FILES: list[tuple[str, str]] = [
    # --- Giải pháp / cấu hình ---
    ("TravelGuide.sln", "File solution Visual Studio (nếu có), gom các project."),
    ("TravelGuide/TravelGuide.csproj", "Cấu hình project MAUI: target framework, package NuGet, asset, font."),
    ("TravelGuide.AdminWeb/TravelGuide.AdminWeb.csproj", "Cấu hình ASP.NET Minimal API + static web."),
    ("TravelGuide/Properties/launchSettings.json", "Cấu hình chạy debug MAUI (profile, env)."),
    ("TravelGuide.AdminWeb/Properties/launchSettings.json", "URL/port khi F5 Admin Web."),
    ("TravelGuide.AdminWeb/appsettings.json", "Cấu hình runtime Admin Web (logging, URLs)."),
    ("TravelGuide.AdminWeb/appsettings.Development.json", "Ghi đè cấu hình môi trường Development."),
    # --- MAUI — điểm vào & shell ---
    ("TravelGuide/MauiProgram.cs", "Đăng ký DI, font, localization, AddAudio, singleton Database/Geofence/Narration, transient các trang."),
    ("TravelGuide/App.xaml", "Resource dictionary toàn app: merge Colors + Styles."),
    ("TravelGuide/App.xaml.cs", "Khởi động: seed POI, messenger GPS (legacy geofence TTS), OnSleep dừng thuyết minh."),
    ("TravelGuide/AppShell.xaml", "Shell: route mặc định tới MainPage."),
    ("TravelGuide/AppShell.xaml.cs", "Đăng ký route MapPage, AudioPage."),
    ("TravelGuide/AppLanguage.cs", "Quản lý mã ngôn ngữ (vi/en/ja/ko/zh), Preferences, culture thread."),
    ("TravelGuide/MapboxConfig.cs", "Đọc token Mapbox: Preferences mapbox_access_token, fallback biến MAPBOX_ACCESS_TOKEN."),
    # --- MAUI — dịch vụ lõi ---
    ("TravelGuide/DatabaseService.cs", "Tải/cache danh sách POI từ API public + fallback extra_places.json; dịch MyMemory khi thiếu field."),
    ("TravelGuide/GeofenceEngine.cs", "So vị trí user với POI (debounce/cooldown); ưu tiên Priority khi vùng chồng; gọi NarrationEngine."),
    ("TravelGuide/GpsBackgroundService.cs", "Vòng lặp GPS định kỳ, gửi LocationMessage + feed GeofenceEngine."),
    ("TravelGuide/NarrationEngine.cs", "Hàng đợi: phát AudioUrl (HTTP) qua Plugin.Maui.Audio, không thì TTS đa ngôn ngữ."),
    ("TravelGuide/TranslationService.cs", "Dịch POI qua MyMemory, cập nhật field đa ngôn ngữ + cache."),
    # --- MAUI — model ---
    ("TravelGuide/Models/TouristPlace.cs", "Entity POI: tọa độ, bán kính, Priority, ảnh, AudioUrl, MapLink, tên/mô tả đa ngôn ngữ."),
    ("TravelGuide/Models/LocationMessage.cs", "Tin nhắn MVVM chứa Location + metadata cho WeakReferenceMessenger."),
    # --- MAUI — trang & UI ---
    ("TravelGuide/MainPage.xaml", "Giao diện chọn ngôn ngữ, tiền tệ, nút vào Home."),
    ("TravelGuide/MainPage.xaml.cs", "Logic cài đặt ban đầu, đồng bộ ResX, điều hướng HomePage."),
    ("TravelGuide/HomePage.xaml", "Dashboard: danh sách POI, tìm kiếm, nút Map/Audio/Gần đây."),
    ("TravelGuide/HomePage.xaml.cs", "Load POI, lọc, mở chi tiết / map / audio, nearby theo GPS."),
    ("TravelGuide/MapPage.xaml", "WebView bản đồ, thanh GPS, banner POI gần, nút reload."),
    ("TravelGuide/MapPage.xaml.cs", "HTML Mapbox hoặc Leaflet/OSM; bridge app://; highlight POI geofence."),
    ("TravelGuide/AudioPage.xaml", "Danh sách POI phát TTS/audio, play all, shuffle, now playing."),
    ("TravelGuide/AudioPage.xaml.cs", "ObservableCollection AudioItem, điều khiển queue NarrationEngine."),
    ("TravelGuide/PlaceDetailPage.xaml", "Chi tiết POI: ảnh, mô tả, phát, nút mở Map link."),
    ("TravelGuide/PlaceDetailPage.xaml.cs", "Refresh theo ngôn ngữ; SpeakAsync; Launcher mở MapLink."),
    ("TravelGuide/MiniPlayerView.xaml", "Thanh mini: tên POI đang phát, stop/next."),
    ("TravelGuide/MiniPlayerView.xaml.cs", "Subscribe sự kiện NarrationEngine, cập nhật UI."),
    # --- MAUI — tài nguyên ---
    ("TravelGuide/Resources/Styles/Colors.xaml", "Bảng màu dùng chung (DynamicResource)."),
    ("TravelGuide/Resources/Styles/Styles.xaml", "Style mặc định ContentPage và control MAUI."),
    ("TravelGuide/Resources/AppResources.resx", "Chuỗi đa ngôn ngữ (ResX) — file gốc bản dịch UI."),
    ("TravelGuide/Resources/AppResources.Designer.cs", "Mã C# auto-generated từ AppResources.resx (không sửa tay)."),
    ("TravelGuide/Resources/Raw/extra_places.json", "Fallback POI nhúng package khi không gọi được API."),
    # --- MAUI — platform ---
    ("TravelGuide/Platforms/Android/MainApplication.cs", "Entry Android: CreateMauiApp → MauiProgram."),
    ("TravelGuide/Platforms/Android/MainActivity.cs", "Activity chính; khởi động Foreground LocationService khi có quyền."),
    ("TravelGuide/Platforms/Android/LocationService.cs", "Service nền Android: GPS định kỳ, messenger + geofence."),
    ("TravelGuide/Platforms/iOS/AppDelegate.cs", "Entry iOS: CreateMauiApp."),
    ("TravelGuide/Platforms/iOS/Program.cs", "Main iOS → UIApplication với AppDelegate."),
    ("TravelGuide/Platforms/MacCatalyst/AppDelegate.cs", "Entry Mac Catalyst: CreateMauiApp."),
    ("TravelGuide/Platforms/MacCatalyst/Program.cs", "Main Mac Catalyst."),
    ("TravelGuide/Platforms/Windows/App.xaml", "Host WinUI cho MAUI Windows."),
    ("TravelGuide/Platforms/Windows/App.xaml.cs", "MauiWinUIApplication, CreateMauiApp."),
    ("TravelGuide/Platforms/Tizen/Main.cs", "Entry Tizen (nếu bật target)."),
    # --- Admin Web ---
    ("TravelGuide.AdminWeb/Program.cs", "Minimal API: login, CRUD POI, public POI, duyệt/từ chối, audio list, tài khoản, dịch, export JSON."),
    ("TravelGuide.AdminWeb/GlobalUsings.cs", "Global using cho namespace Auth, Data, Models, Services."),
    ("TravelGuide.AdminWeb/Data/TravelGuideDb.cs", "SQLite: bảng Poi + UserAccount, seed, CRUD, dịch tự động MyMemory."),
    ("TravelGuide.AdminWeb/Models/Dtos.cs", "Record DTO: PoiDto, Login, UserAccount, ExportPoi, v.v."),
    ("TravelGuide.AdminWeb/Auth/AuthHelper.cs", "Đọc Bearer token từ header, tra AuthStore."),
    ("TravelGuide.AdminWeb/Auth/AuthStore.cs", "Lưu token → principal trong RAM."),
    ("TravelGuide.AdminWeb/Services/PasswordTools.cs", "Hash mật khẩu SHA256 (đăng nhập)."),
    ("TravelGuide.AdminWeb/Services/MyMemoryTranslator.cs", "Gọi API MyMemory dịch vi→đích."),
    ("TravelGuide.AdminWeb/wwwroot/index.html", "Giao diện CMS: login, tab POI / dịch / tài khoản, form Priority & Map link."),
    ("TravelGuide.AdminWeb/wwwroot/app.js", "Logic SPA: fetch API, bảng POI, lưu form, duyệt/từ chối, export."),
    ("TravelGuide.AdminWeb/wwwroot/styles.css", "CSS layout sidebar, card, form, bảng."),
    ("TravelGuide.AdminWeb/scripts/build_code_map_docx.py", "Script tạo Word ánh xạ hàm Admin theo sequence diagram."),
    ("TravelGuide.AdminWeb/scripts/build_file_overview_docx.py", "Script tạo file Word này (tổng quan file)."),
    # --- Sơ đồ / tài liệu ---
    ("Diagrams/USECASE.drawio", "Sơ đồ use case (Draw.io) — tài liệu kiến trúc/đồ án."),
    ("TravelGuide.AdminWeb/AdminWeb_MaNguon_HamVaFile.docx", "Word ánh xạ mã nguồn Admin (sinh bởi script; có thể cập nhật sau khi refactor)."),
]


def main():
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("TravelGuide — Tổng quan từng file").bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Thư mục dự án: {REPO_ROOT}")
    doc.add_paragraph(
        "Bảng dưới liệt kê đường dẫn tương đối (từ thư mục chứa TravelGuide.sln / folder gốc repo) "
        "và mô tả ngắn vai trò file. Một số file phụ (obj, bin) không liệt kê."
    )

    doc.add_heading("Danh sách file", level=1)
    table = doc.add_table(rows=1 + len(FILES), cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "File (đường dẫn tương đối)"
    table.rows[0].cells[1].text = "Tổng quan (làm gì)"
    for i, (path, desc) in enumerate(FILES, start=1):
        table.rows[i].cells[0].text = path
        table.rows[i].cells[1].text = desc

    doc.add_paragraph()
    note = doc.add_paragraph()
    note.add_run("Ghi chú: ").bold = True
    note.add_run(
        "Chạy lại script sau khi thêm/xóa file lớn để cập nhật Word. "
        "Token Mapbox và mật khẩu không lưu trong tài liệu này."
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"Wrote: {OUT}")


if __name__ == "__main__":
    main()
